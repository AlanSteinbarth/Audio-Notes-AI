# =============================================================================
# APLIKACJA AUDIO NOTES AI
# =============================================================================
"""
Kompleksowa aplikacja do zarządzania notatkami głosowymi z wykorzystaniem AI.
*** ENTERPRISE VERSION 2.1.0 - UNIWERSALNA WERSJA ***

AUTOR: Alan Steinbarth
EMAIL: alan.steinbarth@gmail.com  
GITHUB: https://github.com/AlanSteinbarth
WERSJA: 2.1.0
DATA: 2025-06-14

FUNKCJONALNOŚCI:
- Nagrywanie notatek głosowych za pomocą mikrofonu
- Automatyczna transkrypcja audio za pomocą OpenAI Whisper
- Zapisywanie notatek w bazie wektorowej Qdrant z embeddingami
- Semantyczne wyszukiwanie notatek na podstawie treści
- Eksport notatek do formatów TXT, PDF, DOCX
- Edycja i usuwanie istniejących notatek
- Interfejs użytkownika w języku polskim

TECHNOLOGIE:
- Streamlit: Framework dla interfejsu webowego
- OpenAI API: Transkrypcja (Whisper) i embeddingi (text-embedding-3-large)
- Qdrant: Baza danych wektorowych do przechowywania embeddingów
- FPDF2/python-docx: Generowanie dokumentów PDF i DOCX
- audiorecorder: Komponent do nagrywania audio w Streamlit

STRUKTURA PLIKU:
1. Importy i konfiguracja podstawowa
2. Funkcje obsługi API (OpenAI, Qdrant)
3. Funkcje obsługi bazy danych
4. Funkcje eksportu dokumentów
5. Główna funkcja aplikacji z interfejsem użytkownika
"""

# =============================================================================
# IMPORTY I KONFIGURACJA
# =============================================================================

from io import BytesIO
import streamlit as st
import locale
import logging
import io
import os
import platform
from dotenv import dotenv_values
from hashlib import md5
from openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct, Distance, VectorParams
from datetime import datetime
from docx import Document
from typing import Optional
from pathlib import Path

# Importy opcjonalne - tylko flagi, bez komunikatów Streamlit
AUDIORECORDER_AVAILABLE = True
FPDF_AVAILABLE = True
try:
    from audiorecorder import audiorecorder  # type: ignore
except (ImportError, Exception) as e:
    AUDIORECORDER_AVAILABLE = False
    audiorecorder = None
    print(f"⚠️ Audiorecorder nie jest dostępny: {e}")
    print("💡 Spróbuj reinstalację: pip uninstall streamlit-audiorecorder && pip install streamlit-audiorecorder")
try:
    from fpdf import FPDF  # type: ignore
except ImportError:
    FPDF_AVAILABLE = False
    FPDF = None

# Dodaj import OpenAIError jeśli openai jest dostępne
try:
    from openai.error import OpenAIError  # type: ignore
except ImportError:
    OpenAIError = Exception

# =============================================================================
# SPRAWDZENIE ZALEŻNOŚCI SYSTEMOWYCH
# (tylko flagi, bez komunikatów Streamlit)
MISSING_DEPS = []
def check_system_dependencies():
    missing = []
    if os.system("ffmpeg -version > /dev/null 2>&1") != 0:
        missing.append("ffmpeg")
    if os.system("git --version > /dev/null 2>&1") != 0:
        missing.append("git")
    return missing
MISSING_DEPS = check_system_dependencies()
SYSTEM = platform.system()

# =============================================================================
# KONFIGURACJA LOGOWANIA I OBSŁUGA BŁĘDÓW
# =============================================================================

# Konfiguracja logowania do pliku z rotacją i formatowaniem
log_path = Path("app.log").resolve()
logging.basicConfig(
    filename=str(log_path),
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('AudioNotatki')

def log_error(e: Exception, context: Optional[str] = None):
    """
    Centralna funkcja do logowania błędów z wyświetlaniem w interfejsie.
    
    Args:
        e (Exception): Wyjątek do zalogowania
        context (str, optional): Dodatkowy kontekst opisujący gdzie wystąpił błąd
    """
    error_msg = f"{str(e)}"
    if context:
        error_msg = f"{context}: {error_msg}"
    logger.error(error_msg, exc_info=True)
    st.error(error_msg)

# =============================================================================
# KONFIGURACJA LOKALIZACJI I ZMIENNYCH ŚRODOWISKOWYCH
# =============================================================================

# Ustawienie lokalizacji na polską (dla formatowania dat i liczb)
try:
    locale.setlocale(locale.LC_ALL, 'pl_PL.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'Polish_Poland.1250')
    except locale.Error:
        pass  # Kontynuuj bez polskiego locale

# Wczytanie konfiguracji z pliku .env
env = dotenv_values(".env")

# Walidacja obecności wszystkich wymaganych zmiennych środowiskowych
required_env_vars = ["QDRANT_URL", "QDRANT_API_KEY"]  # usunięto OPENAI_API_KEY z wymagań
missing_vars = [var for var in required_env_vars if var not in env or not env[var]]
if missing_vars:
    st.error(f"Brakuje wymaganych zmiennych środowiskowych: {', '.join(missing_vars)}")
    st.info("Skopiuj plik .env.example do .env i uzupełnij wymagane wartości.")
    st.stop()

# =============================================================================
# STAŁE KONFIGURACYJNE
# =============================================================================

# Konfiguracja modeli OpenAI
EMBEDDING_MODEL = "text-embedding-3-large"  # Model do generowania embeddingów tekstu
EMBEDDING_DIM = 3072                         # Wymiarowość wektorów embeddingów
AUDIO_TRANSCRIBE_MODEL = "whisper-1"         # Model do transkrypcji audio

# Konfiguracja bazy danych Qdrant
QDRANT_COLLECTION_NAME = "notes"             # Nazwa kolekcji w bazie wektorowej

# =============================================================================
# FUNKCJE OBSŁUGI API I KLIENTÓW
# =============================================================================

def get_openai_client():
    """Tworzy i zwraca klienta OpenAI z kluczem API z zmiennych środowiskowych."""
    return OpenAI(api_key=env["OPENAI_API_KEY"])

def transcribe_audio(audio_bytes):
    """
    Transkrypcja pliku audio za pomocą OpenAI Whisper API.
    
    Args:
        audio_bytes: Surowe dane audio w formacie WAV
        
    Returns:
        str: Tekst transkrypcji audio lub None w przypadku błędu
    """
    try:
        with st.spinner("Transkrypcja audio..."):
            if not audio_bytes:
                st.error("Nie otrzymano danych audio")
                return None
            openai_client = get_openai_client()
            audio_file = BytesIO(audio_bytes)
            audio_file.name = "audio.mp3"
            transcript = openai_client.audio.transcriptions.create(
                file=audio_file,
                model=AUDIO_TRANSCRIBE_MODEL,
                response_format="text",
            )
            logger.info("Transkrypcja audio zakończona pomyślnie")
            return str(transcript)
    except (ValueError, TypeError, KeyError, ConnectionError) as e:
        log_error(e, "Błąd transkrypcji")
        return None

@st.cache_resource
def get_qdrant_client():
    """
    Tworzy i zwraca klienta bazy danych Qdrant z cache'owaniem.
    
    Returns:
        QdrantClient: Skonfigurowany klient bazy danych Qdrant
    """
    try:
        client = QdrantClient(
            url=env["QDRANT_URL"],
            api_key=env["QDRANT_API_KEY"]
        )
        logger.info("Pomyślnie połączono z Qdrant")
        return client
    except (ConnectionError, ValueError, KeyError) as e:
        log_error(e, "Nie można połączyć się z bazą danych Qdrant")
        st.stop()

def initialize_collection():
    """
    Inicjalizuje kolekcję w bazie Qdrant z odpowiednią konfiguracją wektorów.
    
    Returns:
        QdrantClient: Klient z zainicjalizowaną kolekcją
    """
    try:
        client = get_qdrant_client()
        collections = client.get_collections().collections
        exists = any(collection.name == QDRANT_COLLECTION_NAME for collection in collections)
        if not exists:
            client.create_collection(
                collection_name=QDRANT_COLLECTION_NAME,
                vectors_config=VectorParams(size=EMBEDDING_DIM, distance=Distance.COSINE)
            )
            logger.info("Utworzono nową kolekcję: %s", QDRANT_COLLECTION_NAME)
        return client
    except (ValueError, KeyError, ConnectionError) as e:
        log_error(e, "Błąd podczas inicjalizacji kolekcji Qdrant")
        st.stop()

def verify_openai_key(api_key: str) -> bool:
    """Weryfikuje poprawność klucza OpenAI przez próbę pobrania własnych usage lub modelu."""
    try:
        client = OpenAI(api_key=api_key)
        client.models.list()
        return True
    except (OpenAIError, ValueError, KeyError, ConnectionError):
        return False

# =============================================================================
# FUNKCJE OBSŁUGI BAZY DANYCH I EMBEDDINGÓW
# =============================================================================

def get_embeddings(text: str) -> list[float]:
    """
    Generuje wektor embeddings dla podanego tekstu przy użyciu OpenAI API.
    
    Args:
        text (str): Tekst do przekonwertowania na embedding
        
    Returns:
        list[float]: Lista liczb reprezentująca wektor embeddings lub pusta lista w przypadku błędu
    """
    try:
        openai_client = get_openai_client()
        result = openai_client.embeddings.create(
            input=[text],
            model=EMBEDDING_MODEL,
            dimensions=EMBEDDING_DIM,
        )
        return result.data[0].embedding
    except (ValueError, TypeError, KeyError, ConnectionError) as e:
        log_error(e, "Błąd podczas generowania wektora embeddings")
        return []

def add_note_to_db(note_text, note_id=None):
    """
    Dodaje nową notatkę lub aktualizuje istniejącą w bazie danych Qdrant.
    
    Args:
        note_text (str): Treść notatki do zapisania
        note_id (int, optional): ID notatki (dla aktualizacji) lub None (dla nowej notatki)
        
    Returns:
        int: ID zapisanej notatki
    """
    try:
        qdrant_client = get_qdrant_client()
        if note_id is None:
            points_count = qdrant_client.count(
                collection_name=QDRANT_COLLECTION_NAME,
                exact=True,
            )
            note_id = points_count.count + 1
            created_at = datetime.now().isoformat()
        else:
            created_at = None
        title = generate_note_title(note_text)
        vector = get_embeddings(text=note_text)
        payload = {
            "text": note_text,
            "title": title,
        }
        if created_at:
            payload["created_at"] = created_at
        qdrant_client.upsert(
            collection_name=QDRANT_COLLECTION_NAME,
            points=[
                PointStruct(
                    id=note_id,
                    vector=vector,
                    payload=payload,
                )
            ]
        )
        return note_id
    except Exception as e:  # noqa: E722
        st.error(f"Wystąpił błąd podczas zapisywania notatki: {str(e)}")
        logger.exception("Błąd podczas zapisywania notatki")
        raise

def delete_note_from_db(note_id):
    """
    Usuwa notatkę o podanym ID z bazy danych Qdrant.
    
    Args:
        note_id (int): Unikalny identyfikator notatki do usunięcia
    """
    try:
        qdrant_client = get_qdrant_client()
        from qdrant_client.models import PointIdsList
        qdrant_client.delete(collection_name=QDRANT_COLLECTION_NAME, points_selector=PointIdsList(points=[note_id]))
        st.toast("Notatka usunięta", icon="🗑️")
    except (ConnectionError, ValueError, KeyError) as e:
        st.error(f"Błąd podczas usuwania notatki: {str(e)}")
        logger.exception("Błąd podczas usuwania notatki")

def list_notes_from_db(query=None):
    """
    Pobiera listę notatek z bazy danych z opcjonalnym wyszukiwaniem semantycznym.
    
    Args:
        query (str, optional): Tekst zapytania do wyszukiwania semantycznego
        
    Returns:
        list[dict]: Lista słowników z danymi notatek
    """
    try:
        qdrant_client = get_qdrant_client()
        if not query:
            notes = qdrant_client.scroll(collection_name=QDRANT_COLLECTION_NAME, limit=20)[0]
            result = []
            for note in notes:
                if note.payload and "text" in note.payload:
                    result.append({
                        "id": note.id,
                        "title": note.payload.get("title", "Brak tytułu"),
                        "text": note.payload["text"],
                        "created_at": note.payload.get("created_at", "brak daty"),
                        "score": None,
                    })
            return result
        else:
            query_vector = get_embeddings(text=query)
            notes = qdrant_client.search(
                collection_name=QDRANT_COLLECTION_NAME,
                query_vector=query_vector,
                limit=20,
            )
            result = []
            for note in notes:
                if note.payload and "text" in note.payload:
                    result.append({
                        "id": note.id,
                        "title": note.payload.get("title", "Brak tytułu"),
                        "text": note.payload["text"],
                        "created_at": note.payload.get("created_at", "brak daty"),
                        "score": round(note.score, 3),
                    })
            return result
    except (ConnectionError, ValueError, KeyError) as e:
        st.error(f"Wystąpił błąd podczas pobierania notatek: {str(e)}")
        logger.exception("Błąd podczas pobierania notatek")
        return []

def generate_note_title(note_text):
    """
    Generuje krótki, opisowy tytuł dla notatki przy użyciu OpenAI GPT-3.5.
    
    Args:
        note_text (str): Treść notatki do przeanalizowania
        
    Returns:
        str: Wygenerowany tytuł notatki lub "Brak tytułu" w przypadku błędu
    """
    try:
        with st.spinner("Generowanie tytułu..."):
            openai_client = get_openai_client()
            response = openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system",
                        "content": "Jesteś asystentem, który tworzy krótkie, zwięzłe tytuły na podstawie treści notatek. Tytuł powinien mieć maksymalnie 5 słów."
                    },
                    {
                        "role": "user",
                        "content": f"Wygeneruj krótki tytuł dla tej notatki: {note_text}"
                    }
                ],
                max_tokens=50
            )
            content = None
            if response and hasattr(response, "choices") and response.choices:
                content = response.choices[0].message.content
            if content:
                return content.strip()
            return "Brak tytułu"
    except (ValueError, TypeError, KeyError, ConnectionError) as e:
        log_error(e, "Błąd podczas generowania tytułu")
        return "Brak tytułu"

# =============================================================================
# GŁÓWNA FUNKCJA APLIKACJI I INTERFEJS UŻYTKOWNIKA
# =============================================================================

def main():
    """Główna funkcja aplikacji Streamlit zawierająca cały interfejs użytkownika."""
    # Konfiguracja strony Streamlit z tytułem i layoutem
    st.set_page_config(page_title="🎤 Audio Notes AI 🤖", layout="centered")

    # Sidebar: pole do podania klucza OpenAI
    st.sidebar.header("🔑 Ustawienia API")
    env_key = env.get("OPENAI_API_KEY", "")
    api_key = st.sidebar.text_input("Podaj OpenAI API Key", value=env_key, type="password")
    if not api_key:
        st.error("Podaj klucz OpenAI API w sidebarze lub pliku .env")
        st.stop()
    if not verify_openai_key(api_key):
        st.error("Nieprawidłowy klucz OpenAI API! Sprawdź i wprowadź poprawny klucz w sidebarze lub pliku .env.")
        st.stop()
    # Ustawienie klucza do dalszego użycia
    env["OPENAI_API_KEY"] = api_key

    # Inicjalizacja stanu sesji dla przechowywania danych między interakcjami
    if "note_audio_bytes_md5" not in st.session_state:
        st.session_state["note_audio_bytes_md5"] = None
    if "note_audio_bytes" not in st.session_state:
        st.session_state["note_audio_bytes"] = None
    if "note_text" not in st.session_state:
        st.session_state["note_text"] = ""
    if "note_audio_text" not in st.session_state:
        st.session_state["note_audio_text"] = ""

    # Nagłówek główny aplikacji
    st.title("🎤 Audio Notes AI 🤖")
    st.write("Nagraj notatkę głosową lub wyszukaj w swoich notatkach")
    
    # Inicjalizacja połączenia z bazą danych Qdrant
    try:
        initialize_collection()
    except (ValueError, KeyError, ConnectionError) as e:
        log_error(e)
        st.stop()

    # Komunikaty o zależnościach systemowych
    if MISSING_DEPS:
        msg = f"Brakuje zależności systemowych: {', '.join(MISSING_DEPS)}.\n"
        if SYSTEM == "Darwin":
            msg += "Zainstaluj brakujące pakiety na macOS: brew install " + ' '.join(MISSING_DEPS)
        elif SYSTEM == "Linux":
            msg += "Zainstaluj brakujące pakiety na Linux: sudo apt install " + ' '.join(MISSING_DEPS)
        elif SYSTEM == "Windows":
            msg += "Zainstaluj brakujące pakiety na Windows (np. przez Chocolatey): choco install " + ' '.join(MISSING_DEPS)
        else:
            msg += "Zainstaluj brakujące pakiety odpowiednio dla swojego systemu."
        st.warning(msg)
    # Komunikaty o bibliotekach opcjonalnych
    if not AUDIORECORDER_AVAILABLE:
        st.info("Nagrywanie audio jest niedostępne. Zainstaluj streamlit-audiorecorder: pip install streamlit-audiorecorder")
    if not FPDF_AVAILABLE:
        st.info("Eksport PDF niedostępny. Zainstaluj fpdf: pip install fpdf")

    # Utworzenie trzech głównych zakładek interfejsu użytkownika
    add_tab, search_tab, list_tab = st.tabs(["Dodaj notatkę", "Wyszukaj notatkę", "Lista notatek"])

    # =========================================================================
    # ZAKŁADKA 1: DODAWANIE NOTATEK
    # =========================================================================
    with add_tab:
        if audiorecorder is None:
            st.info("Nagrywanie audio jest niedostępne. Zainstaluj streamlit-audiorecorder.")
        else:
            # Komponent do nagrywania audio z konfigurowalnymi komunikatami
            note_audio = audiorecorder(
                start_prompt="Nagraj notatkę",
                stop_prompt="Zatrzymaj nagrywanie",
            )
            
            # Obsługa nagrania audio - konwersja i przechowywanie w session state
            if note_audio:
                # Eksport nagrania do formatu MP3 w pamięci
                audio = BytesIO()
                note_audio.export(audio, format="mp3")
                st.session_state["note_audio_bytes"] = audio.getvalue()
                
                # Sprawdzenie czy nagranie się zmieniło (hash MD5)
                current_md5 = md5(st.session_state["note_audio_bytes"]).hexdigest()
                if st.session_state["note_audio_bytes_md5"] != current_md5:
                    # Resetowanie poprzednich transkrypcji przy nowym nagraniu
                    st.session_state["note_audio_text"] = ""
                    st.session_state["note_text"] = ""
                    st.session_state["note_audio_bytes_md5"] = current_md5

                # Wyświetlenie odtwarzacza audio
                st.audio(st.session_state["note_audio_bytes"], format="audio/mp3")

                # Przycisk do uruchomienia transkrypcji przez OpenAI Whisper
                if st.button("Transkrybuj audio"):
                    st.session_state["note_audio_text"] = transcribe_audio(st.session_state["note_audio_bytes"])

                # Edytor tekstu do modyfikacji transkrybowanej notatki
                if st.session_state["note_audio_text"]:
                    st.session_state["note_text"] = st.text_area("Edytuj notatkę", value=st.session_state["note_audio_text"])

                # Przycisk zapisu notatki z walidacją długości tekstu
                if st.session_state["note_text"] and st.button("Zapisz notatkę", disabled=not st.session_state["note_text"]):
                    if len(st.session_state["note_text"].strip()) < 5:
                        st.error("Notatka musi mieć co najmniej 5 znaków.")
                    else:
                        add_note_to_db(note_text=st.session_state["note_text"])
                        st.toast("Notatka zapisana", icon="🎉")

    # =========================================================================
    # ZAKŁADKA 2: WYSZUKIWANIE SEMANTYCZNE NOTATEK
    # =========================================================================
    with search_tab:
        # Pole tekstowe do wprowadzenia zapytania wyszukiwania
        query = st.text_input("Wyszukaj notatkę")
        
        # Przycisk uruchamiający wyszukiwanie semantyczne
        if st.button("Szukaj"):
            # Wywołanie funkcji wyszukiwania z parametrem query
            results = list_notes_from_db(query)
            
            # Obsługa przypadku braku wyników
            if not results:
                st.info("Nie znaleziono pasujących notatek.")
            else:
                # Wyświetlenie znalezionych notatek z oceną podobieństwa
                for note in results:
                    with st.container(border=True):
                        st.markdown(f"### {note['title']}")
                        st.markdown(note["text"])
                        # Wyświetlenie oceny podobieństwa jeśli dostępna
                        if note["score"]:
                            st.markdown(f':violet[Podobieństwo: {note["score"]}]')

    # =========================================================================
    # ZAKŁADKA 3: LISTA WSZYSTKICH NOTATEK I ZARZĄDZANIE
    # =========================================================================
    with list_tab:
        st.subheader("Wszystkie notatki")
        
        # Pobranie wszystkich notatek z bazy danych
        notes = list_notes_from_db()
        
        # Obsługa przypadku pustej bazy danych
        if not notes:
            st.info("Brak notatek w bazie.")
        else:
            # Iteracja przez wszystkie notatki z opcjami zarządzania
            for note in notes:
                with st.container(border=True):
                    # Wyświetlenie tytułu, treści i daty utworzenia
                    st.markdown(f"### {note['title']}")
                    st.markdown(note["text"])
                    st.caption(f"Dodano: {note['created_at']}")
                    
                    # Trzy kolumny z przyciskami akcji
                    col1, col2, col3 = st.columns([1,1,2])
                    
                    # Kolumna 1: Przycisk usuwania notatki
                    with col1:
                        if st.button("Usuń", key=f"del_{note['id']}"):
                            delete_note_from_db(note['id'])
                            st.rerun()
                    
                    # Kolumna 2: Przycisk edycji notatki
                    with col2:
                        if st.button("Edytuj", key=f"edit_{note['id']}"):
                            # Przygotowanie danych do edycji w session state
                            st.session_state["edit_note_id"] = note['id']
                            st.session_state["edit_note_text"] = note['text']
                            st.session_state["edit_note_title"] = note['title']
                            st.rerun()
                    
                    # Kolumna 3: Opcje eksportu w różnych formatach
                    with col3:
                        st.download_button(
                            "Eksport TXT",
                            note["text"],
                            file_name=f"notatka_{note['id']}.txt",
                            key=f"txt_{note['id']}"
                        )
                        if FPDF is not None:
                            try:
                                pdf = FPDF()
                                pdf.add_page()
                                pdf.set_font("helvetica", size=12)
                                safe_title = note["title"].encode('ascii', 'ignore').decode('ascii')
                                safe_text = note["text"].encode('ascii', 'ignore').decode('ascii')
                                if not safe_title.strip():
                                    safe_title = f"Notatka {note['id']}"
                                if not safe_text.strip():
                                    safe_text = "Treść zawiera znaki specjalne nieobsługiwane przez PDF"
                                pdf.multi_cell(0, 10, safe_title + "\n\n" + safe_text)
                                pdf_bytes = io.BytesIO()
                                pdf_output = pdf.output()
                                if isinstance(pdf_output, str):
                                    pdf_bytes.write(pdf_output.encode('latin1'))
                                else:
                                    pdf_bytes.write(pdf_output)
                                pdf_bytes.seek(0)
                                st.download_button(
                                    "Eksport PDF",
                                    data=pdf_bytes,
                                    file_name=f"notatka_{note['id']}.pdf",
                                    key=f"pdf_{note['id']}"
                                )
                            except (ValueError, TypeError, KeyError, ConnectionError) as e:
                                logger.error("Błąd podczas generowania PDF: %s", str(e))
                                st.error("Nie udało się wygenerować pliku PDF. Spróbuj eksportu DOCX lub TXT.")
                        else:
                            st.info("Eksport PDF niedostępny. Zainstaluj fpdf.")
                        
                        # EKSPORT DOCX - pełne wsparcie dla polskich znaków
                        try:
                            doc = Document()
                            doc.add_heading(note["title"], 0)
                            doc.add_paragraph(note["text"])
                            docx_bytes = io.BytesIO()
                            doc.save(docx_bytes)
                            docx_bytes.seek(0)
                            st.download_button(
                                "Eksport DOCX",
                                data=docx_bytes,
                                file_name=f"notatka_{note['id']}.docx",
                                key=f"docx_{note['id']}"
                            )
                        except (ValueError, TypeError, KeyError, ConnectionError) as e:
                            logger.error("Błąd podczas generowania DOCX: %s", str(e))
                            st.error("Nie udało się wygenerować pliku DOCX.")

    # =========================================================================
    # TRYB EDYCJI NOTATEK
    # =========================================================================
    if "edit_note_id" in st.session_state:
        with st.form("edit_note_form", clear_on_submit=True):
            st.subheader("Edytuj notatkę")
            
            # Pole tekstowe z aktualną treścią notatki do edycji
            new_text = st.text_area("Treść notatki", value=st.session_state["edit_note_text"])
            
            # Przycisk zapisania zmian z walidacją
            col1, col2 = st.columns([1, 1])
            
            with col1:
                if st.form_submit_button("Zapisz zmiany"):
                    if not new_text or len(new_text.strip()) < 5:
                        st.error("Notatka musi mieć co najmniej 5 znaków.")
                    else:
                        try:
                            add_note_to_db(note_text=new_text, note_id=st.session_state["edit_note_id"])
                            st.toast("Notatka zaktualizowana", icon="✅")
                            # Usunięcie stanu edycji po zapisaniu
                            for key in ["edit_note_id", "edit_note_text", "edit_note_title"]:
                                if key in st.session_state:
                                    del st.session_state[key]
                            st.rerun()
                        except (ValueError, TypeError, KeyError, ConnectionError) as e:
                            st.error(f"Błąd podczas zapisywania: {str(e)}")
            
            with col2:
                if st.form_submit_button("Anuluj"):
                    # Usunięcie stanu edycji bez zapisywania
                    for key in ["edit_note_id", "edit_note_text", "edit_note_title"]:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.rerun()

if __name__ == "__main__":
    main()
